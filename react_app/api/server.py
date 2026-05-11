import os
import sys
import json
import tempfile
from io import BytesIO
from pathlib import Path

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel
from typing import Any, List, Dict

PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

try:
    from dotenv import load_dotenv
except ImportError:
    def load_dotenv(*_args, **_kwargs):
        return False

load_dotenv(PROJECT_ROOT / ".env")
load_dotenv(PROJECT_ROOT / "react_app" / ".env")

from model.model import compute_final_score, compute_score_breakdown
from preprocessing.preprocessing_pipeline import build_feature_vector
from utils.file_extractor import extract_docx_text, extract_image_text, extract_pdf_text
from utils.llm_helper import (
    _make_groq_client,
    chat_with_groq,
    generate_feedback,
)

JOB_OPTIONS = {
    "Frontend Developer": "Looking for a frontend developer with React, JavaScript, Redux, and API integration experience",
    "Backend Developer": "Looking for backend developer with Node.js, Express, databases, and API development",
    "Data Scientist": "Looking for Python, Machine Learning, Pandas, NumPy, and data analysis experience",
    "Full Stack Developer": "Looking for MERN stack developer with React, Node.js, MongoDB and API experience",
    "Custom": "",
}

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".png", ".jpg", ".jpeg"}
RESUME_HINTS = {
    "resume",
    "curriculum vitae",
    "experience",
    "education",
    "skills",
    "projects",
    "work history",
    "employment",
    "internship",
    "certifications",
    "linkedin",
    "github",
    "portfolio",
    "objective",
    "summary",
}

RESOURCE_TOOLS = {
    "resume-checklist": {
        "env": "GROQ_RESUME_CHECKLIST_KEY",
        "title": "Resume checklist",
        "instruction": (
            "Create a practical resume checklist from the analysis. Focus on concrete resume edits, "
            "missing evidence, keyword alignment, metrics, formatting, and what to rewrite first."
        ),
    },
    "skill-gap-map": {
        "env": "GROQ_SKILL_GAP_MAP_KEY",
        "title": "Skill gap map",
        "instruction": (
            "Map the candidate's gaps into skill groups. Identify what is present, what is weak or missing, "
            "priority level, and the fastest practice path for each gap."
        ),
    },
    "project-prompts": {
        "env": "GROQ_PROJECT_PROMPTS_KEY",
        "title": "Project prompts",
        "instruction": (
            "Generate project ideas that prove the missing skills. Each project should include goal, stack, "
            "features, proof points, and resume bullets the candidate can write after building it."
        ),
    },
    "interview-prep": {
        "env": "GROQ_INTERVIEW_PREP_KEY",
        "title": "Interview prep",
        "instruction": (
            "Convert the analysis into interview preparation. Include likely questions, strong talking points, "
            "gap explanations, STAR-style answer outlines, and a short practice plan."
        ),
    },
}

app = FastAPI(title="Resume AI React API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://127.0.0.1:5173", "http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def get_match_status(score):
    if score >= 75:
        return "Strong match"
    if score >= 55:
        return "Moderate match"
    return "Weak match"


def validate_resume_text(resume_text):
    text = (resume_text or "").strip()
    words = [word for word in text.replace("\n", " ").split(" ") if word.strip()]
    lowered = text.lower()
    hint_count = sum(1 for hint in RESUME_HINTS if hint in lowered)
    has_contact_signal = "@" in text or any(char.isdigit() for char in text)

    if len(words) < 35 or hint_count < 1 or not has_contact_signal:
        raise HTTPException(status_code=400, detail="Please upload resume only.")


def normalize_resource_result(parsed, tool_title):
    if not isinstance(parsed, dict):
        parsed = {}

    summary = str(parsed.get("summary") or "").strip()
    if not summary:
        summary = f"{tool_title} generated from the latest resume analysis."

    sections = parsed.get("sections")
    if not isinstance(sections, list) or not sections:
        sections = [
            {
                "title": tool_title,
                "items": [
                    {
                        "label": "Review result",
                        "detail": "The model response did not include structured sections. Re-run this resource or use the dashboard report as the source.",
                        "priority": "Medium",
                    }
                ],
            }
        ]

    normalized_sections = []
    for section in sections:
        if not isinstance(section, dict):
            continue
        title = str(section.get("title") or "Recommendation").strip()
        raw_items = section.get("items")
        if not isinstance(raw_items, list) or not raw_items:
            raw_items = [{"label": title, "detail": str(section.get("detail") or summary), "priority": "Medium"}]
        items = []
        for item in raw_items:
            if isinstance(item, dict):
                items.append(
                    {
                        "label": str(item.get("label") or "Action").strip(),
                        "detail": str(item.get("detail") or item.get("text") or "Review and apply this recommendation.").strip(),
                        "priority": str(item.get("priority") or "Medium").strip(),
                    }
                )
            elif str(item).strip():
                items.append({"label": "Action", "detail": str(item).strip(), "priority": "Medium"})
        normalized_sections.append({"title": title, "items": items})

    next_steps = parsed.get("next_steps")
    if not isinstance(next_steps, list):
        next_steps = []
    next_steps = [str(step).strip() for step in next_steps if str(step).strip()]
    if not next_steps:
        next_steps = [
            "Apply the highest-priority recommendation.",
            "Update the resume draft in the editor.",
            "Run the analyzer again and compare the score.",
        ]

    return {
        "summary": summary,
        "sections": normalized_sections,
        "next_steps": next_steps[:5],
    }


class ResourceRequest(BaseModel):
    analysis: dict
    job_description: str | None = None
    role_template: str | None = None


class EnhanceTextRequest(BaseModel):
    text: str
    field: str = "resume text"
    context: dict[str, Any] = {}

    class Config:
        extra = "forbid"


class ResumeExportRequest(BaseModel):
    draft: dict[str, Any]
    format: str = "docx"

    class Config:
        extra = "forbid"


def call_groq(resource_id, payload):
    tool = RESOURCE_TOOLS.get(resource_id)
    if not tool:
        raise HTTPException(status_code=404, detail="Resource tool not found.")

    api_key = os.getenv(tool["env"], "").strip()
    if not api_key:
        raise HTTPException(
            status_code=500,
            detail=f"Missing {tool['env']} in .env. Add it locally before running this resource.",
        )

    analysis = payload.analysis or {}
    prompt = f"""
You are powering the "{tool['title']}" page in a resume analysis app.

Task:
{tool['instruction']}

Use this latest analysis as the only source of truth:
Role template: {payload.role_template or analysis.get("role_template") or "Unknown"}
Match score: {analysis.get("score")}
Match status: {analysis.get("status")}
Semantic similarity: {analysis.get("similarity")}
Depth gap: {analysis.get("depth_gap")}
Resume signals: {analysis.get("resume_signals")}
Job description signals: {analysis.get("jd_signals")}
Resume preview:
{analysis.get("resume_preview", "")}

API feedback:
{analysis.get("feedback", "")}

Job description:
{payload.job_description or ""}

Return valid JSON only, no markdown fences, with this shape:
{{
  "summary": "one concise overview sentence",
  "sections": [
    {{
      "title": "section title",
      "items": [
        {{
          "label": "short label",
          "detail": "specific recommendation",
          "priority": "High|Medium|Low"
        }}
      ]
    }}
  ],
  "next_steps": ["short action", "short action", "short action"]
}}
""".strip()

    try:
        import groq
    except ImportError as error:
        raise HTTPException(
            status_code=500,
            detail="`groq` package is not installed. Run `pip install groq`.",
        ) from error

    try:
        client = _make_groq_client(groq, api_key)
        response = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": "Return concise, structured, valid JSON for a resume improvement UI.",
                },
                {"role": "user", "content": prompt},
            ],
            model=os.getenv("GROQ_RESOURCE_MODEL", os.getenv("GROQ_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct")),
            temperature=0.35,
            max_completion_tokens=int(os.getenv("GROQ_RESOURCE_MAX_COMPLETION_TOKENS", os.getenv("GROQ_MAX_COMPLETION_TOKENS", "1200"))),
            response_format={"type": "json_object"},
        )
        content = response.choices[0].message.content or "{}"
    except Exception as error:
        raise HTTPException(status_code=502, detail=f"Groq API request failed. {str(error)}") from error

    try:
        parsed = json.loads(content)
    except json.JSONDecodeError:
        parsed = {
            "summary": "The model returned a text response instead of structured JSON.",
            "sections": [
                {
                    "title": tool["title"],
                    "items": [{"label": "Raw response", "detail": content, "priority": "Medium"}],
                }
            ],
            "next_steps": ["Review the raw response", "Re-run the resource if needed"],
        }

    return {
        "resource": tool["title"],
        "result": normalize_resource_result(parsed, tool["title"]),
    }


@app.get("/api/health")
def health():
    return {"ok": True}


@app.get("/api/roles")
def roles():
    return {"roles": JOB_OPTIONS}


@app.post("/api/resources/{resource_id}")
def generate_resource(resource_id: str, payload: ResourceRequest):
    return call_groq(resource_id, payload)


def improve_resume_text_locally(text, field):
    raw_lines = [line.strip(" -•\t") for line in str(text or "").splitlines() if line.strip()]
    if not raw_lines:
        return ""

    weak_starts = ("responsible for", "worked on", "helped with", "involved in", "handled")
    action_verbs = ["Built", "Led", "Improved", "Delivered", "Optimized", "Created", "Reduced", "Increased"]
    improved = []

    for index, line in enumerate(raw_lines):
        clean = " ".join(line.split()).rstrip(".")
        lowered = clean.lower()
        for weak in weak_starts:
            if lowered.startswith(weak):
                clean = clean[len(weak):].strip(" :-")
                break

        starts_with_action = clean.split(" ", 1)[0] in action_verbs
        if "summary" in field.lower():
            improved.append(clean)
            continue

        if not starts_with_action:
            clean = f"{action_verbs[index % len(action_verbs)]} {clean[:1].lower()}{clean[1:]}"
        if not any(char.isdigit() for char in clean):
            clean = f"{clean}, improving delivery quality and role alignment"
        improved.append(clean + ".")

    if "summary" in field.lower():
        return " ".join(improved)[:520]
    return "\n".join(improved)


@app.post("/api/resume/enhance-text")
def enhance_resume_text(payload: EnhanceTextRequest):
    text = (payload.text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="Add text before generating an improved version.")

    api_key = os.getenv("GROQ_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(status_code=500, detail="GROQ_API_KEY is required for resume text generation.")

    try:
        import groq
        client = _make_groq_client(groq, api_key)
        context = payload.context or {}
        prompt = f"""
Rewrite this resume field so it is concise, ATS-friendly, and achievement-focused.

Field: {payload.field}
Target role: {context.get("target_role", "")}
Target industry: {context.get("target_industry", "")}
Career level: {context.get("career_level", "")}
Market/region: {context.get("market", "")}
Job description keywords: {context.get("keywords", "")}

Rules:
- Preserve truthful meaning. Do not invent employers, dates, degrees, certifications, or fake metrics.
- Use standard resume language, no personal pronouns.
- For bullets, start each line with a strong action verb and follow Context, Action, Result where possible.
- Quantify only when the source text already includes numbers; otherwise use measurable wording without fake numbers.
- Keep summaries to 3-4 short lines max.
- Use past tense for completed work and present tense only for current work.
- Return only the rewritten text, no markdown fence.

Original text:
{text}
""".strip()
        response = client.chat.completions.create(
            messages=[
                {"role": "system", "content": "You rewrite resume text into concise ATS-friendly content."},
                {"role": "user", "content": prompt},
            ],
            model=os.getenv("GROQ_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct"),
            temperature=0.25,
            max_completion_tokens=int(os.getenv("GROQ_MAX_COMPLETION_TOKENS", "700")),
        )
        return {"text": (response.choices[0].message.content or "").strip(), "source": "groq"}
    except Exception as error:
        raise HTTPException(status_code=502, detail=f"Groq API request failed. {str(error)}") from error


def bullet_lines(text):
    return [line.strip(" -•\t") for line in str(text or "").splitlines() if line.strip()]


def get_draft_section_order(draft):
    level = str(((draft.get("basics") or {}).get("careerLevel") or "")).lower()
    if "fresher" in level or "intern" in level or "entry" in level:
        return ["summary", "education", "skills", "projects", "experience", "extras"]
    if "changer" in level:
        return ["summary", "skills", "experience", "projects", "education", "extras"]
    if "executive" in level or "senior" in level:
        return ["summary", "skills", "experience", "projects", "education", "extras"]
    return ["summary", "experience", "skills", "projects", "education", "extras"]


def build_resume_plain_text_server(draft):
    basics = draft.get("basics") or {}
    contact = " | ".join(
        str(basics.get(key) or "").strip()
        for key in ["email", "phone", "location", "links"]
        if str(basics.get(key) or "").strip()
    )
    lines = [basics.get("name") or "Candidate Name", basics.get("title") or "", contact, ""]

    def add_summary():
        lines.extend(["Professional Summary", basics.get("summary") or "", ""])

    def add_skills():
        lines.extend(["Skills", draft.get("skills") or "", ""])

    def add_experience():
        lines.append("Experience")
        for item in draft.get("experience") or []:
            lines.extend([f"{item.get('role', '')} - {item.get('company', '')}".strip(" -"), item.get("period") or ""])
            lines.extend(f"- {line}" for line in bullet_lines(item.get("bullets")))
            lines.append("")

    def add_projects():
        lines.append("Projects")
        for item in draft.get("projects") or []:
            lines.extend([item.get("name") or "", item.get("stack") or ""])
            lines.extend(f"- {line}" for line in bullet_lines(item.get("bullets")))
            lines.append("")

    def add_education():
        lines.append("Education")
        for item in draft.get("education") or []:
            lines.extend([f"{item.get('degree', '')} - {item.get('school', '')}".strip(" -"), item.get("period") or "", item.get("detail") or "", ""])

    def add_extras():
        extras = draft.get("extras") or {}
        if extras.get("certifications"):
            lines.extend(["Certifications", extras.get("certifications"), ""])
        if extras.get("achievements"):
            lines.extend(["Achievements", extras.get("achievements"), ""])

    renderers = {
        "summary": add_summary,
        "skills": add_skills,
        "experience": add_experience,
        "projects": add_projects,
        "education": add_education,
        "extras": add_extras,
    }
    for section in get_draft_section_order(draft):
        renderers[section]()
    return "\n".join(str(line).strip() for line in lines).replace("\n\n\n", "\n\n").strip()


def build_resume_docx(draft):
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as error:
        raise HTTPException(status_code=500, detail="python-docx is not installed.") from error

    basics = draft.get("basics") or {}
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = 1
    run = title.add_run(str(basics.get("name") or "Candidate Name"))
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = 1
    subtitle.add_run(str(basics.get("title") or ""))
    contact = " | ".join(str(basics.get(key) or "").strip() for key in ["email", "phone", "location", "links"] if str(basics.get(key) or "").strip())
    contact_para = document.add_paragraph()
    contact_para.alignment = 1
    contact_para.add_run(contact)

    def heading(text):
        para = document.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(11.5)

    def bullets(text):
        for line in bullet_lines(text):
            document.add_paragraph(line, style="List Bullet")

    def add_summary():
        heading("Professional Summary")
        document.add_paragraph(str(basics.get("summary") or ""))

    def add_skills():
        heading("Skills")
        document.add_paragraph(str(draft.get("skills") or ""))

    def add_experience():
        heading("Experience")
        for item in draft.get("experience") or []:
            para = document.add_paragraph()
            para.add_run(f"{item.get('role', '')} - {item.get('company', '')}".strip(" -")).bold = True
            if item.get("period"):
                para.add_run(f" | {item.get('period')}")
            bullets(item.get("bullets"))

    def add_projects():
        heading("Projects")
        for item in draft.get("projects") or []:
            para = document.add_paragraph()
            para.add_run(str(item.get("name") or "")).bold = True
            if item.get("stack"):
                para.add_run(f" | {item.get('stack')}")
            bullets(item.get("bullets"))

    def add_education():
        heading("Education")
        for item in draft.get("education") or []:
            para = document.add_paragraph()
            para.add_run(f"{item.get('degree', '')} - {item.get('school', '')}".strip(" -")).bold = True
            if item.get("period"):
                para.add_run(f" | {item.get('period')}")
            if item.get("detail"):
                document.add_paragraph(str(item.get("detail")))

    def add_extras():
        extras = draft.get("extras") or {}
        if extras.get("certifications"):
            heading("Certifications")
            document.add_paragraph(str(extras.get("certifications")))
        if extras.get("achievements"):
            heading("Achievements")
            document.add_paragraph(str(extras.get("achievements")))

    renderers = {
        "summary": add_summary,
        "skills": add_skills,
        "experience": add_experience,
        "projects": add_projects,
        "education": add_education,
        "extras": add_extras,
    }
    for section in get_draft_section_order(draft):
        renderers[section]()

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


@app.post("/api/resume/export")
def export_resume(payload: ResumeExportRequest):
    fmt = (payload.format or "docx").lower()
    draft = payload.draft or {}
    name = str(((draft.get("basics") or {}).get("name") or "resume")).lower()
    safe_name = "".join(char if char.isalnum() else "-" for char in name).strip("-") or "resume"

    if fmt == "docx":
        output = build_resume_docx(draft)
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}-resume.docx"'},
        )
    if fmt == "txt":
        return Response(
            build_resume_plain_text_server(draft),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}-resume.txt"'},
        )
    if fmt == "json":
        return Response(
            json.dumps(draft, indent=2),
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}-resume.json"'},
        )
    raise HTTPException(status_code=400, detail="Supported export formats are docx, txt, and json.")


@app.post("/api/analyze")
async def analyze_resume(
    resume: UploadFile = File(...),
    job_description: str = Form(...),
):
    jd = (job_description or "").strip()
    if not jd:
        raise HTTPException(status_code=400, detail="Job description is required.")

    suffix = Path(resume.filename or "").suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise HTTPException(
            status_code=400,
            detail="Please upload resume only.",
        )

    temp_path = None
    try:
        contents = await resume.read()
        if not contents:
            raise HTTPException(status_code=400, detail="Please upload resume only.")

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(contents)
            temp_path = temp_file.name

        if suffix == ".pdf":
            resume_text = extract_pdf_text(temp_path)
        elif suffix == ".docx":
            resume_text = extract_docx_text(temp_path)
        else:
            resume_text = extract_image_text(temp_path)

        validate_resume_text(resume_text)
        features = build_feature_vector(resume_text, jd)
        final_score, similarity, depth_gap = compute_final_score(features)
        score_breakdown = compute_score_breakdown(features)
        feedback = generate_feedback(resume_text, jd, final_score)

        return {
            "score": final_score,
            "status": get_match_status(final_score),
            "similarity": round(float(similarity), 4),
            "depth_gap": depth_gap,
            "score_breakdown": score_breakdown,
            "ats_score": score_breakdown["ats_score"],
            "skill_coverage": score_breakdown["skill_coverage"],
            "keyword_coverage": score_breakdown["keyword_coverage"],
            "evidence_score": score_breakdown["evidence_score"],
            "matched_skills": features.get("matched_skills", []),
            "missing_skills": features.get("missing_skills", []),
            "resume_skills": features.get("resume_skills", []),
            "jd_skills": features.get("jd_skills", []),
            "matched_keywords": features.get("matched_keywords", []),
            "resume_signals": features.get("resume_signals", []),
            "jd_signals": features.get("jd_signals", []),
            "feedback": feedback,
            "resume_preview": (resume_text or "")[:2500],
            "resume_text": (resume_text or "")[:12000],
        }
    finally:
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)


class ChatRequest(BaseModel):
    resume_text: str
    job_description: str
    messages: List[Dict[str, str]]

@app.post("/api/chat")
def chat_endpoint(req: ChatRequest):
    try:
        messages = []
        for msg in (req.messages or [])[-12:]:
            role = msg.get("role")
            content = str(msg.get("content", "")).strip()
            if role in {"user", "assistant"} and content:
                messages.append({"role": role, "content": content})
        if not req.resume_text.strip():
            raise HTTPException(status_code=400, detail="Resume text is required for chat.")
        if not messages:
            raise HTTPException(status_code=400, detail="Add at least one chat message.")

        reply = chat_with_groq(req.resume_text, req.job_description, messages)
        return {"reply": reply}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
