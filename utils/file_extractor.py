import os
import subprocess
import tempfile

# Trigger reload for PyPDF2
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    from PIL import Image
    from PIL import ImageOps
except ImportError:
    Image = None
    ImageOps = None

try:
    from docx import Document
except ImportError:
    Document = None


def extract_pdf_text(file_path):
    if PyPDF2 is None:
        return ""

    text_chunks = []
    try:
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                text_chunks.append(page_text)
        return "\n".join(text_chunks).strip()
    except Exception:
        return ""


def extract_image_text(image_path):
    if Image is None:
        return ""

    try:
        img = Image.open(image_path)
        img = ImageOps.grayscale(img)
        img = ImageOps.autocontrast(img)

        if pytesseract is not None:
            text = pytesseract.image_to_string(img)
        else:
            text = _extract_image_text_with_tesseract_cli(img)
        return text.strip()
    except Exception:
        return ""


def _extract_image_text_with_tesseract_cli(img):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
        temp_path = tmp.name

    try:
        img.save(temp_path)
        completed = subprocess.run(
            ["tesseract", temp_path, "stdout", "--psm", "6"],
            check=False,
            capture_output=True,
            text=True,
            timeout=30,
        )
        if completed.returncode != 0:
            return ""
        return completed.stdout
    except (OSError, subprocess.SubprocessError):
        return ""
    finally:
        if os.path.exists(temp_path):
            os.unlink(temp_path)


def extract_docx_text(file_path):
    if Document is None:
        return ""

    try:
        document = Document(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    value = cell.text.strip()
                    if value:
                        table_cells.append(value)
        return "\n".join([*paragraphs, *table_cells]).strip()
    except Exception:
        return ""
